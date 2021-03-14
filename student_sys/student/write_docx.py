from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

import re
import os, sys
import django

PROJ_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.append(PROJ_ROOT)
os.environ.setdefault("DJANGO_SETTINGS_MODULE", "student_sys.settings")
django.setup()

from student.models import Student, Project, StudentJoinProject, StudentScholar, StudentOrganization

ACTIVITIES = [
    "思想成长类活动",
    "实习实践类活动",
    "志愿公益类活动",
    "团学社团类活动",
    "创新创业类活动",
    "技术特长类活动",
    "文体学术类活动",
]
EX_ACTIVITIES = ["学生组织经历", "所获奖励"]
CONTENT_HEAD = ["序号", "活动日期", "活动名称", "主办单位", "认证状态"]
EX_CONTENT_HEAD1 = ["序号", "开始时间", "结束时间", "组织名称", "职位", "知道单位", "认证状态"]
EX_CONTENT_HEAD2 = ["序号", "评奖学年", "奖学金名称", "教学金级别", "认证状态"]


class ReportDocx:
    def __init__(
        self,
        student_id,
        activities=ACTIVITIES,
        ex_activities=EX_ACTIVITIES,
        content_head=CONTENT_HEAD,
        ex_content_head1=EX_CONTENT_HEAD1,
        ex_content_head2=EX_CONTENT_HEAD2,
    ):
        self.student_id = student_id  # 学号
        self.activities = activities  # 参加的活动类别
        self.ex_activities = ex_activities  # 其它经历
        self.content_head = content_head  # 参加活动类别的表格标题
        self.ex_content_head1 = ex_content_head1  # 组织活动经历的表格标题
        self.ex_content_head2 = ex_content_head2  # 奖学金
        self.student_report = {}
        self.hint_list = []  # 提示信息
        self.document = Document()  # 初始化docx文件
        self.get_report()  # 生成docx文件

    def get_report(self):
        try:
            s = Student.objects.get(student_id=self.student_id)
            if s.grade == 0:
                grade = int("20" + str(s.student_id[:2]))
            else:
                grade = s.grade

            if s.sex == 1:
                sex = "男"
            elif s.sex == 2:
                sex = "女"
            else:
                sex = "未知"

            sps = StudentJoinProject.objects.filter(s_id=s.student_id)
            p_info = {}  # 写入参加活动信息
            for sp in sps:
                p_name = sp.project_name
                try:
                    p = Project.objects.get(name=p_name, semester=sp.semester)
                    category = p.category
                    if category in p_info:
                        c = int(p_info[category][-1][0])
                        p_info[category].append(
                            [
                                str(c + 1),
                                p.semester,
                                p.name,
                                p.department_name,
                                p.certify_state if p.certify_state else "已认证",
                            ]
                        )
                    else:
                        p_info[category] = [
                            ["1", p.semester, p.name, p.department_name, p.certify_state if p.certify_state else "已认证"]
                        ]
                except Exception as e:
                    self.hint_list.append([p_name, str(e)])
                    print(p_name, e)

            self.student_report = {
                "基础信息": [s.department_name, str(grade), s.major_name, s.student_id, s.name, sex],
                "参加活动经历": p_info,
            }
        except Exception as e:
            print(self.student_id, e)
            self.hint_list.append([self.student_id, str(e)])

    def _setCellBackgroundColor(self, cell, hexColor="CDC9C9"):
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls("w"), color_value=hexColor))
        cell._tc.get_or_add_tcPr().append(shading_elm_1)

    def generate_docx(self):
        # 标题行
        h1 = self.document.add_heading(level=1)
        h1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
        c = h1.add_run("上海对外经贸大学学生第二课堂经历证明")
        c.font.size = Pt(16)
        c.font.name = "Times New Roman"
        c._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
        c.font.color.rgb = RGBColor(0, 0, 0)
        c.font.bold = True

        # 表格1：写学生信息
        base_info = self.student_report["基础信息"]
        h2 = self.document.add_heading(level=1)
        data1 = [
            ["所在学院：%s" % base_info[0], "年级：%s" % base_info[1], "专业：%s" % base_info[2]],
            ["学号：%s" % base_info[3], "姓名：%s" % base_info[4], "性别：%s" % base_info[5]],
        ]

        table = self.document.add_table(rows=3, cols=4)
        for i, data in enumerate(data1):
            row = table.rows[i].cells
            for j, d in enumerate(data):
                if j == 0:
                    row[j].width = Cm(6)

                p = row[j].paragraphs[0]
                t = p.add_run(d)
                t.font.size = Pt(10.5)
                t.font.name = "Times New Roman"
                t._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
                t.font.color.rgb = RGBColor(0, 0, 0)
                t.font.bold = True

        # 参加活动详细
        content = self.student_report["参加活动经历"]
        for i, title in enumerate(self.activities):
            h2 = self.document.add_heading("", level=2)
            c = h2.add_run(title)
            c.font.size = Pt(9)
            c.font.name = "Times New Roman"
            c._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            c.font.color.rgb = RGBColor(0, 0, 0)
            c.font.bold = True

            table = self.document.add_table(rows=len(content.get(title, [])) + 1, cols=5)
            head_row = table.rows[0].cells  # 标题行单元格
            for j, cell in enumerate(head_row):
                if j == 1:
                    cell.width = Cm(4)

                self._setCellBackgroundColor(cell)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                p = cell.paragraphs[0]
                # p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                c = p.add_run(self.content_head[j])
                c.font.size = Pt(7.5)
                c.font.name = "Times New Roman"
                c._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                c.font.color.rgb = RGBColor(0, 0, 0)
                c.font.bold = True

            if len(content.get(title, [])) == 0:
                pass
            else:
                for k, d in enumerate(content[title]):  # 内容行单元格
                    row = table.rows[k + 1].cells
                    for l, cell in enumerate(row):
                        if j == 1:
                            cell.width = Cm(4)

                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        p = cell.paragraphs[0]
                        # p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        c = p.add_run(d[l])
                        c.font.size = Pt(7.5)
                        c.font.name = "Times New Roman"
                        c._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                        c.font.color.rgb = RGBColor(0, 0, 0)
                        c.font.bold = False

        for i, title in enumerate(self.ex_activities):  # 数据库无数据的表格
            h2 = self.document.add_heading("", level=2)
            c = h2.add_run(title)
            c.font.size = Pt(9)
            c.font.name = "Times New Roman"
            c._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            c.font.color.rgb = RGBColor(0, 0, 0)
            c.font.bold = True

            table = self.document.add_table(rows=2, cols=5)
            head_row = table.rows[0].cells  # 标题行单元格

            for j, cell in enumerate(head_row):
                if j == 1:
                    cell.width = Cm(4)

                self._setCellBackgroundColor(cell)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                p = cell.paragraphs[0]
                # p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                c = p.add_run(self.ex_content_head[j])
                c.font.size = Pt(7.5)
                c.font.name = "Times New Roman"
                c._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                c.font.color.rgb = RGBColor(0, 0, 0)
                c.font.bold = True

    def download(self):
        self.document.save(f"{self.student_id}-report.docx")


if __name__ == "__main__":
    s_id = "19080010"
    reportDocx = ReportDocx(s_id)
    reportDocx.download()